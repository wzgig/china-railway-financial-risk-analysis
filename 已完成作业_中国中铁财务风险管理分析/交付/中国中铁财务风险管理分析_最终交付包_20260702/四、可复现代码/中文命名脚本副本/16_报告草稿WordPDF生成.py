"""Build the formatted DOCX/PDF draft from the Markdown paper."""

from __future__ import annotations

import re
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt


ROOT = Path(__file__).resolve().parents[1]
DRAFT_MD = ROOT / "paper/draft.md"
REFERENCES_MD = ROOT / "paper/references_gbt7714.md"
TEMPLATE_DOC = ROOT / "长沙理工大学本科毕业设计（论文）撰写规范样张.doc"
BUILD_DIR = ROOT / "outputs/docx"
BUILD_MD = BUILD_DIR / "course_paper_export.md"
REFERENCE_DOCX = BUILD_DIR / "format_reference.docx"
DOCX_OUTPUT = ROOT / "paper/course_paper_draft.docx"
PDF_OUTPUT = ROOT / "paper/course_paper_draft.pdf"

SOFFICE = Path(r"C:\Program Files\LibreOffice\program\soffice.com")


def command_path(command: str) -> str | None:
    found = shutil.which(command)
    return found


def run(command: list[str], cwd: Path | None = None) -> None:
    subprocess.run(command, cwd=cwd or ROOT, check=True)


def clean_reference_text(text: str) -> str:
    lines: list[str] = []
    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if not line:
            lines.append("")
            continue
        if line.startswith("# "):
            continue
        if line.startswith("> "):
            continue
        lines.append(line)
    return "\n".join(lines).strip()


def build_markdown() -> None:
    draft = DRAFT_MD.read_text(encoding="utf-8")
    references = clean_reference_text(REFERENCES_MD.read_text(encoding="utf-8"))
    draft = draft.replace("../docs/assets/figures/", "docs/assets/figures/")
    draft = re.sub(
        r"\n## 参考文献\s*\n\s*见 `paper/references_gbt7714\.md`。\s*$",
        "\n## 参考文献\n\n" + references + "\n",
        draft,
        flags=re.S,
    )
    if not draft.startswith("# "):
        raise ValueError("paper/draft.md should start with a level-1 title")
    title_line, body = draft.split("\n", 1)
    title = title_line.removeprefix("# ").strip()
    metadata = (
        "---\n"
        f"title: \"{title}\"\n"
        "subtitle: \"课程报告草稿\"\n"
        "date: \"2026-06-30\"\n"
        "lang: zh-CN\n"
        "---\n\n"
    )
    BUILD_DIR.mkdir(parents=True, exist_ok=True)
    BUILD_MD.write_text(metadata + body.strip() + "\n", encoding="utf-8")


def convert_template_to_reference_docx() -> Path | None:
    if not TEMPLATE_DOC.exists() or not SOFFICE.exists():
        return None
    BUILD_DIR.mkdir(parents=True, exist_ok=True)
    run(
        [
            str(SOFFICE),
            "--headless",
            "--convert-to",
            "docx",
            "--outdir",
            str(BUILD_DIR),
            str(TEMPLATE_DOC),
        ],
        cwd=ROOT,
    )
    converted = BUILD_DIR / (TEMPLATE_DOC.stem + ".docx")
    if converted.exists():
        shutil.copyfile(converted, REFERENCE_DOCX)
        return REFERENCE_DOCX
    return None


def pandoc_build(reference_docx: Path | None) -> None:
    pandoc = command_path("pandoc")
    if not pandoc:
        raise FileNotFoundError("pandoc was not found")
    command = [
        pandoc,
        str(BUILD_MD),
        "--standalone",
        f"--resource-path={ROOT};{ROOT / 'docs'};{ROOT / 'paper'};{BUILD_DIR}",
        "-o",
        str(DOCX_OUTPUT),
    ]
    if reference_docx and reference_docx.exists():
        command.insert(-2, f"--reference-doc={reference_docx}")
    run(command, cwd=ROOT)


def set_run_font(run, ascii_font: str = "Times New Roman", east_asia_font: str = "宋体") -> None:
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)


def set_style_font(document: Document, style_name: str, size_pt: float, bold: bool = False, east_asia: str = "宋体") -> None:
    style = document.styles[style_name]
    style.font.name = "Times New Roman"
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    field_separate = OxmlElement("w:fldChar")
    field_separate.set(qn("w:fldCharType"), "separate")
    field_text = OxmlElement("w:t")
    field_text.text = "1"
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.append(field_begin)
    run._r.append(instr)
    run._r.append(field_separate)
    run._r.append(field_text)
    run._r.append(field_end)
    set_run_font(run)


def clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def insert_page_break_before(paragraph) -> None:
    new_p = OxmlElement("w:p")
    new_r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    new_r.append(br)
    new_p.append(new_r)
    paragraph._p.addprevious(new_p)


def paragraph_in_table(paragraph) -> bool:
    parent = paragraph._element.getparent()
    while parent is not None:
        if parent.tag == qn("w:tc"):
            return True
        parent = parent.getparent()
    return False


def emu_to_twips(value: int) -> int:
    return int(value / 635)


def set_or_add_property(parent, tag: str):
    for child in parent:
        if child.tag == qn(tag):
            return child
    child = OxmlElement(tag)
    parent.append(child)
    return child


def set_cell_width(cell, width: int) -> None:
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = set_or_add_property(tc_pr, "w:tcW")
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(emu_to_twips(width)))


def fit_table_to_page(table, page_width: int) -> None:
    rows = table.rows
    if not rows:
        return
    col_count = len(rows[0].cells)
    if col_count == 0:
        return

    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = set_or_add_property(tbl_pr, "w:tblW")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(emu_to_twips(page_width)))
    tbl_layout = set_or_add_property(tbl_pr, "w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")

    widths = [int(page_width / col_count)] * col_count
    for row in rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(idx, col_count - 1)])


def polish_docx() -> None:
    document = Document(DOCX_OUTPUT)
    for section in document.sections:
        section.start_type = WD_SECTION.NEW_PAGE
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(25)
        section.bottom_margin = Mm(25)
        section.left_margin = Mm(30)
        section.right_margin = Mm(25)
        for paragraph in section.header.paragraphs:
            clear_paragraph(paragraph)
        if section.footer.paragraphs:
            clear_paragraph(section.footer.paragraphs[0])
            add_page_number(section.footer.paragraphs[0])

    set_style_font(document, "Normal", 10.5, east_asia="宋体")
    for style_name, size, east_asia in [
        ("Title", 18, "黑体"),
        ("Subtitle", 12, "宋体"),
        ("Heading 1", 15, "黑体"),
        ("Heading 2", 14, "黑体"),
        ("Heading 3", 12, "黑体"),
    ]:
        if style_name in document.styles:
            set_style_font(document, style_name, size, bold=True, east_asia=east_asia)

    for idx, paragraph in enumerate(document.paragraphs):
        style_name = paragraph.style.name if paragraph.style else ""
        text = paragraph.text.strip()
        is_caption = bool(re.match(r"^(图|表)\s*\d+", text))
        is_keywords = text.startswith("关键词：") or text.startswith("Key words:")
        is_reference = bool(re.match(r"^\[\d+\]", text))
        is_title_page_line = idx < 3 or style_name in {"Title", "Subtitle", "Date"} or text == "2026-06-30"
        for run in paragraph.runs:
            is_heading = style_name.startswith("Heading") or style_name in {"Title", "Subtitle", "Date"}
            east_asia = "黑体" if style_name.startswith("Heading") or style_name == "Title" or idx == 0 else "宋体"
            set_run_font(run, east_asia_font=east_asia)
            if idx == 0:
                run.font.size = Pt(16)
                run.font.bold = True
            elif idx == 1:
                run.font.size = Pt(12)
                run.font.bold = False
            elif idx == 2:
                run.font.size = Pt(11)
                run.font.bold = False
            elif not is_heading:
                run.font.size = Pt(10.5)
                run.font.bold = False
            elif style_name.startswith("Heading"):
                run.font.bold = True
        fmt = paragraph.paragraph_format
        if is_title_page_line:
            fmt.first_line_indent = None
            fmt.space_before = Pt(90) if idx == 0 else Pt(0)
            fmt.space_after = Pt(16) if idx == 0 else Pt(10)
            fmt.line_spacing = 1.25
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif style_name.startswith("Heading"):
            fmt.first_line_indent = None
            fmt.space_before = Pt(10)
            fmt.space_after = Pt(6)
            fmt.line_spacing = 1.25
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif is_caption:
            fmt.first_line_indent = None
            fmt.space_before = Pt(4)
            fmt.space_after = Pt(6)
            fmt.line_spacing = 1.15
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif is_keywords:
            fmt.first_line_indent = None
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(4)
            fmt.line_spacing = 1.25
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif is_reference:
            fmt.left_indent = Pt(18)
            fmt.first_line_indent = Pt(-18)
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(4)
            fmt.line_spacing = 1.15
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif not paragraph_in_table(paragraph) and paragraph.text.strip():
            fmt.left_indent = None
            fmt.first_line_indent = Pt(21)
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(0)
            fmt.line_spacing = 1.5
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    max_width = document.sections[0].page_width - document.sections[0].left_margin - document.sections[0].right_margin
    for table in document.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        fit_table_to_page(table, max_width)
        try:
            table.style = "Table Grid"
        except KeyError:
            pass
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.first_line_indent = None
                    paragraph.paragraph_format.line_spacing = 1.15
                    for run in paragraph.runs:
                        set_run_font(run, east_asia_font="宋体")
                        run.font.size = Pt(9)

    for shape in document.inline_shapes:
        if shape.width > max_width:
            ratio = max_width / shape.width
            shape.width = int(shape.width * ratio)
            shape.height = int(shape.height * ratio)

    for paragraph in document.paragraphs[:8]:
        is_date_style = paragraph.style and paragraph.style.name == "Date"
        if is_date_style or paragraph.text.strip() == "2026-06-30":
            paragraph.add_run().add_break(WD_BREAK.PAGE)
            break

    document.save(DOCX_OUTPUT)


def convert_docx_to_pdf() -> None:
    if not SOFFICE.exists():
        raise FileNotFoundError("LibreOffice soffice.com was not found")
    run(
        [
            str(SOFFICE),
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(PDF_OUTPUT.parent),
            str(DOCX_OUTPUT),
        ],
        cwd=ROOT,
    )
    produced = PDF_OUTPUT.parent / (DOCX_OUTPUT.stem + ".pdf")
    if produced != PDF_OUTPUT and produced.exists():
        produced.replace(PDF_OUTPUT)
    if not PDF_OUTPUT.exists():
        raise FileNotFoundError(f"Expected {PDF_OUTPUT}")


def main() -> int:
    build_markdown()
    reference_docx = convert_template_to_reference_docx()
    pandoc_build(reference_docx)
    polish_docx()
    convert_docx_to_pdf()
    print(f"wrote {DOCX_OUTPUT}")
    print(f"wrote {PDF_OUTPUT}")
    print(f"wrote {BUILD_MD}")
    if reference_docx:
        print(f"used reference docx {reference_docx}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
