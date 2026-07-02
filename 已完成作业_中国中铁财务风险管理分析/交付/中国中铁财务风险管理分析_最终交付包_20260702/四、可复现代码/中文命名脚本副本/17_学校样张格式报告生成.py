"""Build a Word draft following the Changsha University formatting sample."""

from __future__ import annotations

import re
import shutil
import subprocess
from dataclasses import dataclass
from pathlib import Path

import pdfplumber
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt


ROOT = Path(__file__).resolve().parents[1]
DRAFT_MD = ROOT / "paper/draft.md"
REFERENCES_MD = ROOT / "paper/references_gbt7714.md"
TEMPLATE_DOC = ROOT / "长沙理工大学本科毕业设计（论文）撰写规范样张.doc"
BUILD_DIR = ROOT / "outputs/docx/template_formatted"
TEMPLATE_DOCX = BUILD_DIR / "format_sample.docx"
SCHOOL_LOGO = BUILD_DIR / "media/image2.jpeg"
DOCX_OUTPUT = ROOT / "paper/course_paper_formatted.docx"
PDF_OUTPUT = ROOT / "paper/course_paper_formatted.pdf"
FIRST_PASS_DOCX = BUILD_DIR / "course_paper_formatted_first_pass.docx"
FIRST_PASS_PDF = BUILD_DIR / "course_paper_formatted_first_pass.pdf"
SOFFICE = Path(r"C:\Program Files\LibreOffice\program\soffice.com")

SHORT_HEADER_TITLE = "中国中铁财务风险管理分析"


@dataclass
class Block:
    kind: str
    text: str = ""
    level: int = 0
    image_path: Path | None = None
    caption: str = ""


def run(command: list[str], cwd: Path | None = None) -> None:
    subprocess.run(command, cwd=cwd or ROOT, check=True)


def convert_template_assets() -> None:
    BUILD_DIR.mkdir(parents=True, exist_ok=True)
    if TEMPLATE_DOC.exists() and SOFFICE.exists():
        run(
            [
                str(SOFFICE),
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                str(BUILD_DIR),
                str(TEMPLATE_DOC),
            ],
            cwd=ROOT,
        )
        converted = BUILD_DIR / (TEMPLATE_DOC.stem + ".docx")
        if converted.exists():
            shutil.copyfile(converted, TEMPLATE_DOCX)

    if TEMPLATE_DOCX.exists():
        import zipfile

        media_dir = BUILD_DIR / "media"
        media_dir.mkdir(parents=True, exist_ok=True)
        with zipfile.ZipFile(TEMPLATE_DOCX) as archive:
            for item in archive.namelist():
                if item == "word/media/image2.jpeg":
                    (media_dir / "image2.jpeg").write_bytes(archive.read(item))


def clean_reference_text(text: str) -> str:
    lines: list[str] = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line or line.startswith("# ") or line.startswith("> "):
            continue
        lines.append(line)
    return "\n".join(lines)


def read_markdown() -> str:
    draft = DRAFT_MD.read_text(encoding="utf-8")
    references = clean_reference_text(REFERENCES_MD.read_text(encoding="utf-8"))
    draft = draft.replace("../docs/assets/figures/", "docs/assets/figures/")
    return re.sub(
        r"\n## 参考文献\s*\n\s*见 `paper/references_gbt7714\.md`。\s*$",
        "\n## 参考文献\n\n" + references + "\n",
        draft,
        flags=re.S,
    )


def parse_markdown(markdown: str) -> tuple[str, list[Block]]:
    blocks: list[Block] = []
    title = ""
    buffer: list[str] = []
    in_references = False

    def flush() -> None:
        nonlocal buffer
        if buffer:
            blocks.append(Block(kind="paragraph", text=" ".join(part.strip() for part in buffer).strip()))
            buffer = []

    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            flush()
            continue
        image_match = re.match(r"!\[(?P<caption>.+?)\]\((?P<path>.+?)\)", line)
        if image_match:
            flush()
            blocks.append(
                Block(
                    kind="image",
                    image_path=ROOT / image_match.group("path"),
                    caption=normalize_caption(image_match.group("caption")),
                )
            )
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            flush()
            level = len(heading.group(1))
            text = heading.group(2).strip()
            in_references = text == "参考文献"
            if level == 1:
                title = text
            else:
                blocks.append(Block(kind="heading", text=text, level=level))
            continue
        if in_references and re.match(r"^\[\d+\]", line):
            flush()
            blocks.append(Block(kind="paragraph", text=line))
            continue
        buffer.append(line)
    flush()
    if not title:
        raise ValueError("paper/draft.md should start with a level-1 title")
    return title, blocks


def normalize_caption(text: str) -> str:
    match = re.match(r"图\s*(\d+)\s*(.+)", text.strip())
    if match:
        return f"图{match.group(1)}  {match.group(2).strip()}"
    match = re.match(r"表\s*(\d+)\s*(.+)", text.strip())
    if match:
        return f"表{match.group(1)}  {match.group(2).strip()}"
    return text.strip()


def set_run_font(run, size: float, east_asia: str = "宋体", ascii_font: str = "Times New Roman", bold: bool = False) -> None:
    run.font.name = ascii_font
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)


def set_paragraph_spacing(paragraph, line_spacing: float = 1.5, before: float = 0, after: float = 0) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing = line_spacing
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)


def add_field(paragraph, field_name: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_name
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(text)
    run._r.append(end)
    set_run_font(run, 9, east_asia="宋体")


def add_bottom_border(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_borders = p_pr.find(qn("w:pBdr"))
    if p_borders is None:
        p_borders = OxmlElement("w:pBdr")
        p_pr.append(p_borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_borders.append(bottom)


def reset_page_number(section, start: int = 1) -> None:
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))


def configure_section(section, footer: bool, body_total_pages: int | None = None) -> None:
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(30)
    section.right_margin = Mm(20)
    section.header_distance = Mm(15)
    section.footer_distance = Mm(17.5)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    for p in section.header.paragraphs:
        p.clear()
    header = section.header.paragraphs[0]
    header.paragraph_format.tab_stops.add_tab_stop(Mm(160), WD_TAB_ALIGNMENT.RIGHT)
    if SCHOOL_LOGO.exists():
        header.add_run().add_picture(str(SCHOOL_LOGO), width=Mm(42))
    header.add_run("\t")
    title_run = header.add_run(SHORT_HEADER_TITLE)
    set_run_font(title_run, 9, east_asia="宋体")
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_bottom_border(header)

    for p in section.footer.paragraphs:
        p.clear()
    if footer:
        p = section.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("第 ")
        set_run_font(run, 9, east_asia="宋体")
        add_field(p, "PAGE")
        run = p.add_run(" 页 共 ")
        set_run_font(run, 9, east_asia="宋体")
        if body_total_pages is None:
            add_field(p, "SECTIONPAGES")
        else:
            run = p.add_run(str(body_total_pages))
            set_run_font(run, 9, east_asia="宋体")
        run = p.add_run(" 页")
        set_run_font(run, 9, east_asia="宋体")


def split_front_matter(blocks: list[Block]) -> dict[str, str | list[Block]]:
    data: dict[str, str | list[Block]] = {
        "zh_abstract": "",
        "zh_keywords": "",
        "en_abstract": "",
        "en_keywords": "",
        "body": [],
    }
    current = None
    body_started = False
    body: list[Block] = []
    for block in blocks:
        if block.kind == "heading" and re.match(r"^\d+\s+", block.text):
            body_started = True
        if body_started:
            body.append(block)
            continue
        if block.kind == "heading":
            if block.text == "摘要":
                current = "zh_abstract"
            elif block.text == "Abstract":
                current = "en_abstract"
            continue
        if block.kind == "paragraph":
            if block.text.startswith("关键词："):
                data["zh_keywords"] = block.text.removeprefix("关键词：").strip()
            elif block.text.startswith("Key words:"):
                data["en_keywords"] = block.text.removeprefix("Key words:").strip()
            elif current in {"zh_abstract", "en_abstract"}:
                data[current] = (str(data[current]) + "\n" + block.text).strip()
    data["body"] = body
    return data


def build_toc_entries(body: list[Block]) -> list[tuple[int, str]]:
    entries: list[tuple[int, str]] = []
    for block in body:
        if block.kind != "heading":
            continue
        if block.text == "参考文献":
            entries.append((1, block.text))
        elif re.match(r"^\d+\s+", block.text):
            entries.append((1, block.text))
        elif re.match(r"^\d+\.\d+\s+", block.text):
            entries.append((2, block.text))
        elif re.match(r"^\d+\.\d+\.\d+\s+", block.text):
            entries.append((3, block.text))
    return entries


def add_title_page(document: Document, title: str, front: dict[str, str | list[Block]]) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=28, after=24)
    run = p.add_run(title)
    set_run_font(run, 16, east_asia="黑体", bold=True)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=18)
    run = p.add_run("摘要")
    set_run_font(run, 15, east_asia="黑体", bold=True)

    add_body_paragraph(document, str(front["zh_abstract"]), first_indent=True)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=12)
    run = p.add_run("关键词：")
    set_run_font(run, 14, east_asia="黑体", bold=True)
    run = p.add_run(str(front["zh_keywords"]))
    set_run_font(run, 12, east_asia="宋体")

    document.add_page_break()

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=28, after=24)
    run = p.add_run("FINANCIAL RISK MANAGEMENT ANALYSIS OF CHINA RAILWAY GROUP")
    set_run_font(run, 16, east_asia="Times New Roman", bold=True)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=18)
    run = p.add_run("ABSTRACT")
    set_run_font(run, 15, east_asia="Times New Roman", bold=True)

    add_body_paragraph(document, str(front["en_abstract"]), first_indent=True, east_asia="Times New Roman")

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=12)
    run = p.add_run("Key  words: ")
    set_run_font(run, 14, east_asia="Times New Roman", bold=True)
    run = p.add_run(str(front["en_keywords"]))
    set_run_font(run, 12, east_asia="Times New Roman")


def add_toc(document: Document, entries: list[tuple[int, str]], page_numbers: dict[str, int] | None = None) -> None:
    document.add_page_break()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=12, after=18)
    run = p.add_run("目  录")
    set_run_font(run, 16, east_asia="黑体", bold=True)

    for level, text in entries:
        p = document.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Pt(0 if level == 1 else 24 if level == 2 else 48)
        p.paragraph_format.tab_stops.add_tab_stop(Mm(160), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        run = p.add_run(text)
        set_run_font(run, 12, east_asia="黑体" if level == 1 else "宋体", bold=(level == 1))
        p.add_run("\t")
        page = page_numbers.get(text, "") if page_numbers else ""
        run = p.add_run(str(page))
        set_run_font(run, 12, east_asia="宋体")


def add_body_heading(document: Document, text: str, level: int) -> None:
    p = document.add_paragraph()
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=12, after=12)
        run = p.add_run(text)
        set_run_font(run, 15, east_asia="黑体", bold=True)
    elif level == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(p, before=8, after=8)
        run = p.add_run(text)
        set_run_font(run, 14, east_asia="黑体", bold=True)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(p, before=6, after=6)
        run = p.add_run(text)
        set_run_font(run, 12, east_asia="黑体", bold=True)


def add_body_paragraph(document: Document, text: str, first_indent: bool = True, east_asia: str = "宋体") -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p)
    if first_indent:
        p.paragraph_format.first_line_indent = Pt(24)
    add_text_with_citations(p, text, east_asia=east_asia)


def add_text_with_citations(paragraph, text: str, east_asia: str = "宋体") -> None:
    citation_re = re.compile(r"(\[\d+(?:\s*[-,，]\s*\d+)*\])")
    for part in citation_re.split(text):
        if not part:
            continue
        run = paragraph.add_run(part)
        set_run_font(run, 12, east_asia=east_asia)
        if citation_re.fullmatch(part):
            run.font.superscript = True
            run.font.size = Pt(9)


def add_reference_paragraph(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p)
    p.paragraph_format.left_indent = Pt(21)
    p.paragraph_format.first_line_indent = Pt(-21)
    run = p.add_run(text)
    set_run_font(run, 12, east_asia="宋体")


def add_image_block(document: Document, block: Block) -> None:
    if block.image_path and block.image_path.exists():
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=8, after=4)
        p.add_run().add_picture(str(block.image_path), width=Mm(145))
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=8)
    run = p.add_run(block.caption)
    set_run_font(run, 10.5, east_asia="宋体")


def add_body(document: Document, body: list[Block]) -> None:
    in_references = False
    for block in body:
        if block.kind == "heading":
            if block.text == "参考文献":
                in_references = True
                p = document.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_paragraph_spacing(p, before=12, after=12)
                run = p.add_run("参考文献")
                set_run_font(run, 15, east_asia="黑体", bold=True)
                continue
            if re.match(r"^\d+\s+", block.text):
                add_body_heading(document, block.text, 1)
            elif re.match(r"^\d+\.\d+\s+", block.text):
                add_body_heading(document, block.text, 2)
            else:
                add_body_heading(document, block.text, 3)
        elif block.kind == "image":
            add_image_block(document, block)
        elif block.kind == "paragraph":
            if in_references and block.text.startswith("["):
                add_reference_paragraph(document, block.text)
            else:
                add_body_paragraph(document, block.text)


def build_docx(
    path: Path,
    page_numbers: dict[str, int] | None = None,
    body_total_pages: int | None = None,
) -> None:
    title, blocks = parse_markdown(read_markdown())
    front = split_front_matter(blocks)
    body = front["body"]
    assert isinstance(body, list)
    toc_entries = build_toc_entries(body)

    document = Document()
    configure_section(document.sections[0], footer=False)
    add_title_page(document, title, front)
    add_toc(document, toc_entries, page_numbers=page_numbers)

    body_section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_section(body_section, footer=True, body_total_pages=body_total_pages)
    reset_page_number(body_section, 1)
    add_body(document, body)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    if not SOFFICE.exists():
        raise FileNotFoundError("LibreOffice soffice.com was not found")
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    run(
        [
            str(SOFFICE),
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(pdf_path.parent),
            str(docx_path),
        ],
        cwd=ROOT,
    )
    produced = pdf_path.parent / (docx_path.stem + ".pdf")
    if produced.exists() and produced != pdf_path:
        produced.replace(pdf_path)
    if not pdf_path.exists():
        raise FileNotFoundError(f"Expected {pdf_path}")


def extract_layout_info(pdf_path: Path) -> tuple[dict[str, int], int | None]:
    _, blocks = parse_markdown(read_markdown())
    body = split_front_matter(blocks)["body"]
    assert isinstance(body, list)
    headings = [text for _, text in build_toc_entries(body)]
    pages: dict[str, int] = {}
    body_start = None
    with pdfplumber.open(pdf_path) as pdf:
        page_texts = [(idx + 1, page.extract_text() or "") for idx, page in enumerate(pdf.pages)]
    for abs_page, text in page_texts:
        normalized = re.sub(r"\s+", " ", text)
        if "1 引言" in normalized and "基础设施建设行业" in normalized:
            body_start = abs_page
            break
    if body_start is None:
        return pages, None
    for heading in headings:
        target = re.sub(r"\s+", " ", heading)
        for abs_page, text in page_texts:
            if abs_page < body_start:
                continue
            normalized = re.sub(r"\s+", " ", text)
            if target in normalized:
                pages[heading] = abs_page - body_start + 1
                break
    return pages, len(page_texts) - body_start + 1


def main() -> int:
    convert_template_assets()
    build_docx(FIRST_PASS_DOCX)
    convert_docx_to_pdf(FIRST_PASS_DOCX, FIRST_PASS_PDF)
    page_numbers, body_total_pages = extract_layout_info(FIRST_PASS_PDF)
    build_docx(DOCX_OUTPUT, page_numbers=page_numbers, body_total_pages=body_total_pages)
    convert_docx_to_pdf(DOCX_OUTPUT, PDF_OUTPUT)
    print(f"wrote {DOCX_OUTPUT}")
    print(f"wrote {PDF_OUTPUT}")
    print(f"toc entries with pages: {len(page_numbers)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
