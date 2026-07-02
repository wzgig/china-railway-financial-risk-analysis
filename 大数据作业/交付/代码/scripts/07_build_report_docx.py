# -*- coding: utf-8 -*-
"""
用途：以中国中铁财务风险管理分析报告最终版的 Word 样式为模板，
     生成大数据作业分析报告 DOCX。
输入：paper/协鑫能科算电协同价值创造网络分析报告.md，
     已完成作业_中国中铁财务风险管理分析/交付/.../中国中铁财务风险管理分析报告（格式化备份版）.docx，
     outputs/figures/*.png
输出：paper/协鑫能科算电协同价值创造网络分析报告.docx，
     outputs/docx/协鑫能科算电协同价值创造网络分析报告.docx，
     logs/07_build_report_docx.log
说明：中铁报告文件只作为只读样式模板，不改动归档目录。
"""

from __future__ import annotations

import re
import shutil
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
WORKSPACE_ROOT = PROJECT_ROOT.parent
REPORT_MD = PROJECT_ROOT / "paper" / "协鑫能科算电协同价值创造网络分析报告.md"
PAPER_DOCX = PROJECT_ROOT / "paper" / "协鑫能科算电协同价值创造网络分析报告.docx"
OUTPUT_DOCX = PROJECT_ROOT / "outputs" / "docx" / "协鑫能科算电协同价值创造网络分析报告.docx"
LOG_PATH = PROJECT_ROOT / "logs" / "07_build_report_docx.log"

TEMPLATE_DOCX = (
    WORKSPACE_ROOT
    / "已完成作业_中国中铁财务风险管理分析"
    / "交付"
    / "中国中铁财务风险管理分析_最终交付包_20260702"
    / "一、最终报告"
    / "中国中铁财务风险管理分析报告（格式化备份版）.docx"
)
FALLBACK_TEMPLATE = WORKSPACE_ROOT / "已完成作业_中国中铁财务风险管理分析" / "paper" / "course_paper_formatted.docx"

FIGURES = {
    "overview": (
        PROJECT_ROOT / "outputs" / "figures" / "value_network_overview.png",
        "图1  算电协同价值创造网络总图",
    ),
    "centrality": (
        PROJECT_ROOT / "outputs" / "figures" / "centrality_top_nodes.png",
        "图2  PageRank 中心性排名图",
    ),
    "community": (
        PROJECT_ROOT / "outputs" / "figures" / "community_network.png",
        "图3  社群结构网络图",
    ),
    "stage": (
        PROJECT_ROOT / "outputs" / "figures" / "stage_evolution.png",
        "图4  2024-2027 年阶段演化图",
    ),
    "cost": (
        PROJECT_ROOT / "outputs" / "figures" / "cost_mechanism_mapping.png",
        "图5  成本机制与边类型映射图",
    ),
}


@dataclass
class Block:
    kind: str
    text: str
    level: int = 0


def clean_inline(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = text.replace("->", "→")
    return text.strip()


def parse_report() -> tuple[str, str, str, list[Block]]:
    if not REPORT_MD.exists():
        raise FileNotFoundError(f"缺少报告 Markdown：{REPORT_MD.relative_to(PROJECT_ROOT)}")
    lines = REPORT_MD.read_text(encoding="utf-8").splitlines()
    title = ""
    abstract_parts: list[str] = []
    keywords = ""
    blocks: list[Block] = []
    current_heading = ""
    buffer: list[str] = []
    in_abstract = False
    skip_abstract_heading = False

    def flush() -> None:
        nonlocal buffer
        if buffer:
            text = clean_inline(" ".join(part.strip() for part in buffer if part.strip()))
            if text:
                if in_abstract and not blocks:
                    abstract_parts.append(text)
                else:
                    blocks.append(Block(kind="paragraph", text=text))
            buffer = []

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            flush()
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            flush()
            level = len(heading.group(1))
            text = clean_inline(heading.group(2))
            if level == 1:
                title = text
                continue
            current_heading = text
            if text == "摘要":
                in_abstract = True
                skip_abstract_heading = True
                continue
            in_abstract = False
            skip_abstract_heading = False
            if text == "参考资料":
                text = "参考文献"
            blocks.append(Block(kind="heading", text=text, level=level))
            continue

        if skip_abstract_heading and line.startswith("关键词："):
            flush()
            keywords = clean_inline(line.removeprefix("关键词："))
            in_abstract = False
            skip_abstract_heading = False
            continue

        if current_heading == "参考资料" and re.match(r"^\[\d+\]", line):
            flush()
            blocks.append(Block(kind="reference", text=clean_inline(line)))
            continue

        buffer.append(line)

    flush()
    abstract = "\n".join(abstract_parts).strip()
    if not title:
        raise ValueError("报告 Markdown 缺少一级标题")
    if not abstract:
        raise ValueError("报告 Markdown 缺少摘要内容")
    return title, abstract, keywords, blocks


def choose_template() -> Path:
    if TEMPLATE_DOCX.exists():
        return TEMPLATE_DOCX
    if FALLBACK_TEMPLATE.exists():
        return FALLBACK_TEMPLATE
    raise FileNotFoundError("找不到中国中铁最终版 DOCX 模板或备份模板")


def clear_document_body(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_run_font(run, size: float, east_asia: str = "宋体", ascii_font: str = "Times New Roman", bold: bool = False) -> None:
    run.font.name = ascii_font
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)


def set_style_font(document: Document, style_name: str, size: float, east_asia: str, bold: bool = False) -> None:
    if style_name not in document.styles:
        return
    style = document.styles[style_name]
    style.font.name = "Times New Roman"
    style.font.size = Pt(size)
    style.font.bold = bold
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)


def configure_document(document: Document) -> None:
    for section in document.sections:
        section.start_type = WD_SECTION.NEW_PAGE
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(25)
        section.bottom_margin = Mm(25)
        section.left_margin = Mm(30)
        section.right_margin = Mm(25)
        section.header_distance = Mm(15)
        section.footer_distance = Mm(17.5)

    set_style_font(document, "Normal", 12, "宋体")
    set_style_font(document, "Title", 16, "黑体", bold=True)
    set_style_font(document, "Heading 1", 15, "黑体", bold=True)
    set_style_font(document, "Heading 2", 14, "黑体", bold=True)
    set_style_font(document, "Heading 3", 12, "黑体", bold=True)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(text)
    run._r.append(end)
    set_run_font(run, 9)


def configure_header_footer(document: Document) -> None:
    for section in document.sections:
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        for paragraph in section.header.paragraphs:
            paragraph.clear()
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run("协鑫能科算电协同价值创造网络分析")
        set_run_font(run, 9, east_asia="宋体")

        for paragraph in section.footer.paragraphs:
            paragraph.clear()
        add_page_number(section.footer.paragraphs[0])


def add_heading(document: Document, text: str, level: int) -> None:
    p = document.add_paragraph()
    if level <= 2:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(text)
        set_run_font(run, 15, east_asia="黑体", bold=True)
    elif level == 3:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        set_run_font(run, 13, east_asia="黑体", bold=True)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        set_run_font(run, 12, east_asia="黑体", bold=True)


def add_paragraph(document: Document, text: str, *, first_indent: bool = True) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if first_indent:
        p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text)
    set_run_font(run, 12, east_asia="宋体")


def add_reference(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Pt(21)
    p.paragraph_format.first_line_indent = Pt(-21)
    run = p.add_run(text)
    set_run_font(run, 10.5, east_asia="宋体")


def add_figure(document: Document, key: str) -> None:
    path, caption = FIGURES[key]
    if not path.exists():
        return
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(str(path), width=Mm(145))

    caption_p = document.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.space_after = Pt(8)
    run = caption_p.add_run(caption)
    set_run_font(run, 10.5, east_asia="宋体")


def add_title_page(document: Document, title: str, abstract: str, keywords: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(48)
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run(title)
    set_run_font(run, 16, east_asia="黑体", bold=True)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("摘  要")
    set_run_font(run, 15, east_asia="黑体", bold=True)

    add_paragraph(document, abstract)

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run("关键词：")
    set_run_font(run, 12, east_asia="黑体", bold=True)
    run = p.add_run(keywords)
    set_run_font(run, 12, east_asia="宋体")

    document.add_page_break()


def maybe_insert_figure(document: Document, text: str) -> None:
    if text.startswith("整体网络是一个有向加权网络"):
        add_figure(document, "overview")
    elif text.startswith("从中心性指标看"):
        add_figure(document, "centrality")
    elif text.startswith("社群结构可以分为"):
        add_figure(document, "community")
    elif text.startswith("2027 年累计节点达到"):
        add_figure(document, "stage")
    elif text.startswith("从成本角度看"):
        add_figure(document, "cost")


def build_docx() -> Path:
    title, abstract, keywords, blocks = parse_report()
    template_path = choose_template()
    try:
        template_display = template_path.relative_to(WORKSPACE_ROOT)
    except ValueError:
        template_display = Path(template_path.name)
    document = Document(template_path)
    clear_document_body(document)
    configure_document(document)
    configure_header_footer(document)

    add_title_page(document, title, abstract, keywords)
    for block in blocks:
        if block.kind == "heading":
            add_heading(document, block.text, block.level)
        elif block.kind == "reference":
            add_reference(document, block.text)
        else:
            add_paragraph(document, block.text)
            maybe_insert_figure(document, block.text)

    PAPER_DOCX.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(PAPER_DOCX)
    shutil.copy2(PAPER_DOCX, OUTPUT_DOCX)

    LOG_PATH.parent.mkdir(parents=True, exist_ok=True)
    LOG_PATH.write_text(
        "\n".join(
            [
                f"运行时间：{datetime.now():%Y-%m-%d %H:%M:%S}",
                f"模板文件：{template_display}",
                f"输出文件：{PAPER_DOCX.relative_to(PROJECT_ROOT)}",
                f"备份文件：{OUTPUT_DOCX.relative_to(PROJECT_ROOT)}",
                f"段落块数量：{len(blocks)}",
            ]
        )
        + "\n",
        encoding="utf-8",
    )
    return PAPER_DOCX


def main() -> None:
    output = build_docx()
    print(f"已生成：{output.relative_to(PROJECT_ROOT)}")
    print(f"已同步：{OUTPUT_DOCX.relative_to(PROJECT_ROOT)}")


if __name__ == "__main__":
    main()
